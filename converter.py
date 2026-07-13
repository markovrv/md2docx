import re
from io import BytesIO
from typing import Dict, Optional

from bs4 import BeautifulSoup, NavigableString, Tag
import markdown as md_lib
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

from omml import latex_to_omml

FORMULA_PLACEHOLDER = re.compile(r'%%FORMULA_(DISPLAY|INLINE)_(\d+)%%')
INLINE_DOLLAR = re.compile(r'(?<!\$)\$(?!\$)([^$\n]+?)(?<!\$)\$(?!\$)')
DISPLAY_DOLLAR = re.compile(r'\$\$(.+?)\$\$', re.DOTALL)
INLINE_PAREN = re.compile(r'\\\((.*?)\\\)')
DISPLAY_BRACKET = re.compile(r'\\\[(.*?)\\\]', re.DOTALL)


def _render_latex_fallback(formula: str, display: bool = False) -> Optional[bytes]:
    try:
        fig = plt.figure(figsize=(8, 0.8))
        fig.patch.set_alpha(0)
        fig.text(0.5, 0.5, f'${formula}$', fontsize=14,
                 ha='center', va='center', fontfamily='DejaVu Sans')
        buf = BytesIO()
        fig.savefig(buf, format='png', dpi=120,
                    bbox_inches='tight', pad_inches=0.05, transparent=True)
        plt.close(fig)
        buf.seek(0)
        return buf.read()
    except Exception as e:
        print(f"LaTeX fallback img error: {formula} -> {e}")
        plt.close('all')
        return None


class MarkdownToDocx:
    def __init__(self):
        self.doc: Optional[Document] = None
        self.formula_omml: Dict[str, etree.Element] = {}
        self.formula_images: Dict[str, bytes] = {}
        self._list_counters: Dict[int, int] = {}

    def convert(self, md_text: str) -> BytesIO:
        self.doc = Document()
        self._setup_styles()
        self.formula_omml = {}
        self.formula_images = {}
        self._list_counters = {}

        md_text = self._extract_formulas(md_text)
        md_text = self._preprocess_strikethrough(md_text)
        html = md_lib.markdown(md_text, extensions=['extra'])

        soup = BeautifulSoup(html, 'html.parser')
        body = soup.body if soup.body else soup
        self._process_children(self.doc, body)

        buf = BytesIO()
        self.doc.save(buf)
        buf.seek(0)
        return buf

    def _setup_styles(self):
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15

    def _preprocess_strikethrough(self, text: str) -> str:
        return re.sub(r'~~(.+?)~~', r'<del>\1</del>', text)

    def _extract_formulas(self, text: str) -> str:
        counter = [0]
        saved: Dict[str, str] = {}

        def _protect(m, prefix):
            key = f'%%{prefix}{len(saved)}%%'
            saved[key] = m.group(0)
            return key

        text = re.sub(r'```.*?```', lambda m: _protect(m, 'CB'), text, flags=re.DOTALL)
        text = re.sub(r'`[^`\n]+`', lambda m: _protect(m, 'IC'), text)

        def repl_display(m):
            formula = m.group(1).strip()
            ph = f'%%FORMULA_DISPLAY_{counter[0]}%%'
            omml = latex_to_omml(formula, True)
            if omml is not None:
                self.formula_omml[ph] = omml
            else:
                img = _render_latex_fallback(formula, True)
                if img:
                    self.formula_images[ph] = img
            counter[0] += 1
            return '\n\n' + ph + '\n\n'

        def repl_inline(m):
            formula = m.group(1).strip()
            ph = f'%%FORMULA_INLINE_{counter[0]}%%'
            omml = latex_to_omml(formula, False)
            if omml is not None:
                self.formula_omml[ph] = omml
            else:
                img = _render_latex_fallback(formula, False)
                if img:
                    self.formula_images[ph] = img
            counter[0] += 1
            return ph

        text = DISPLAY_DOLLAR.sub(repl_display, text)
        text = DISPLAY_BRACKET.sub(repl_display, text)
        text = INLINE_DOLLAR.sub(repl_inline, text)
        text = INLINE_PAREN.sub(repl_inline, text)

        for key, value in saved.items():
            text = text.replace(key, value)

        return text

    def _process_children(self, parent, container):
        for child in list(container.children):
            if isinstance(child, NavigableString):
                text = str(child).strip()
                if text:
                    p = parent.add_paragraph()
                    self._process_inline(p, text, {})
                continue
            if not isinstance(child, Tag):
                continue

            tag = child.name

            if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
                level = int(tag[1])
                p = parent.add_heading('', level=level)
                self._process_inline(p, child, {})

            elif tag == 'p':
                self._process_para(parent, child)

            elif tag in ('ul', 'ol'):
                self._process_list(parent, child)

            elif tag == 'pre':
                self._process_code(parent, child)

            elif tag == 'blockquote':
                self._process_blockquote(parent, child)

            elif tag == 'table':
                self._process_table(parent, child)

            elif tag == 'hr':
                p = parent.add_paragraph()
                pPr = p.paragraph_format.element.get_or_add_pPr()
                pBdr = parse_xml(
                    f'<w:pBdr {nsdecls("w")}>'
                    f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="AAAAAA"/>'
                    f'</w:pBdr>'
                )
                pPr.append(pBdr)

            elif tag in ('div', 'body', 'html', 'section', 'article', 'main'):
                self._process_children(parent, child)

            elif tag == 'img':
                self._insert_image(parent, child)

            elif tag == 'dl':
                pass

            else:
                self._process_children(parent, child)

    def _process_para(self, parent, tag):
        children = list(tag.children)
        if len(children) == 1 and isinstance(children[0], Tag) and children[0].name == 'img':
            self._insert_image(parent, children[0])
            return
        p = parent.add_paragraph()
        self._process_inline(p, tag, {})

    def _process_inline(self, container, element, fmt):
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                self._emit_text(container, text, fmt)
                continue
            if not isinstance(child, Tag):
                continue

            nfmt = dict(fmt)

            if child.name in ('strong', 'b'):
                nfmt['bold'] = True
                self._process_inline(container, child, nfmt)
            elif child.name in ('em', 'i'):
                nfmt['italic'] = True
                self._process_inline(container, child, nfmt)
            elif child.name == 'code':
                nfmt['code'] = True
                self._emit_text(container, child.get_text(), nfmt)
            elif child.name in ('del', 's', 'strike'):
                nfmt['strikethrough'] = True
                self._process_inline(container, child, nfmt)
            elif child.name == 'a':
                nfmt['link'] = child.get('href', '')
                self._process_inline(container, child, nfmt)
            elif child.name == 'img':
                self._insert_image(container, child)
            elif child.name == 'br':
                container.add_run('\n')
            elif child.name == 'ins':
                nfmt['underline'] = True
                self._process_inline(container, child, nfmt)
            else:
                self._process_inline(container, child, nfmt)

    def _add_run(self, container, text, fmt):
        run = container.add_run(text)
        if fmt.get('bold'):
            run.bold = True
        if fmt.get('italic'):
            run.italic = True
        if fmt.get('code'):
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
        if fmt.get('strikethrough'):
            run.font.strike = True
        if fmt.get('underline'):
            run.font.underline = True
        if fmt.get('link'):
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.font.underline = True

    def _emit_text(self, container, text, fmt):
        last_end = 0
        for m in FORMULA_PLACEHOLDER.finditer(text):
            before = text[last_end:m.start()]
            if before:
                self._add_run(container, before, fmt)
            self._put_formula_omml(container, m.group(0))
            last_end = m.end()
        remaining = text[last_end:]
        if remaining:
            self._add_run(container, remaining, fmt)

    def _put_formula_omml(self, container, placeholder):
        omml = self.formula_omml.get(placeholder)
        is_display = 'DISPLAY' in placeholder

        if omml is not None:
            if is_display:
                for r in list(container._element):
                    if r.tag.endswith('}r'):
                        container._element.remove(r)
                container._element.append(omml)
                container.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                container._element.append(omml)
        else:
            self._put_formula_fallback_image(container, placeholder, is_display)

    def _put_formula_fallback_image(self, container, placeholder, is_display):
        img_data = self.formula_images.get(placeholder)
        if not img_data:
            ptext = placeholder.replace('%%FORMULA_INLINE_', '$').replace('%%FORMULA_DISPLAY_', '$$').replace('%%', '$$').replace('_INLINE_', '$').replace('_DISPLAY_', '$$')
            cleaner = re.sub(r'_(\d+)', '', ptext)
            self._add_run(container, cleaner, {'code': True})
            return
        run = container.add_run()
        try:
            from PIL import Image as PILImage
            pil = PILImage.open(BytesIO(img_data))
            w, h = pil.size
            target_h = 40 if is_display else 20
            scale = target_h / h if h else 1
            run.add_picture(BytesIO(img_data), width=Pt(w * scale))
        except Exception:
            try:
                run.add_picture(BytesIO(img_data), width=Inches(3))
            except Exception:
                pass
        if is_display and hasattr(container, 'paragraph_format'):
            container.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _insert_image(self, parent, tag):
        src = tag.get('src', '')
        if not src or 'FORMULA_' in src:
            return
        try:
            if src.startswith(('http://', 'https://')):
                import requests
                resp = requests.get(src, timeout=10)
                resp.raise_for_status()
                data = resp.content
            elif src.startswith('data:'):
                import base64
                _, encoded = src.split(',', 1)
                data = base64.b64decode(encoded)
            else:
                with open(src, 'rb') as f:
                    data = f.read()
            run = parent.add_run()
            run.add_picture(BytesIO(data), width=Inches(4))
        except Exception as e:
            print(f"Image error: {src} -> {e}")
            parent.add_run('[Image]')

    def _process_list(self, parent, tag, level=0):
        ordered = tag.name == 'ol'
        self._list_counters.setdefault(level, 1)

        indent = level * 0.35

        for item in tag.find_all('li', recursive=False):
            nested = [c for c in item.children if isinstance(c, Tag) and c.name in ('ul', 'ol')]
            text_items = [c for c in item.children if not (isinstance(c, Tag) and c.name in ('ul', 'ol'))]

            has_para = any(isinstance(c, Tag) and c.name == 'p' for c in text_items)

            if not has_para:
                p = parent.add_paragraph()
                p.paragraph_format.left_indent = Inches(indent + 0.25)
                p.paragraph_format.first_line_indent = Inches(-0.25)

                bullet = self._bullet(ordered, level)
                run = p.add_run(bullet)
                run.font.size = Pt(11)

                for c in text_items:
                    self._process_inline(p, c, {}) if isinstance(c, Tag) else self._emit_text(p, str(c), {})
            else:
                for c in text_items:
                    if isinstance(c, Tag) and c.name == 'p':
                        p = parent.add_paragraph()
                        p.paragraph_format.left_indent = Inches(indent + 0.25)
                        p.paragraph_format.first_line_indent = Inches(-0.25)
                        bullet = self._bullet(ordered, level)
                        run = p.add_run(bullet)
                        run.font.size = Pt(11)
                        self._process_inline(p, c, {})
                    elif isinstance(c, NavigableString):
                        txt = str(c).strip()
                        if txt:
                            p = parent.add_paragraph()
                            p.paragraph_format.left_indent = Inches(indent + 0.25)
                            run = p.add_run(self._bullet(ordered, level) + txt)
                            run.font.size = Pt(11)

            for n in nested:
                if ordered:
                    self._list_counters[level + 1] = 1
                self._process_list(parent, n, level + 1)

        if level in self._list_counters:
            self._list_counters[level] = 1

    def _bullet(self, ordered, level):
        if ordered:
            n = self._list_counters.get(level, 1)
            self._list_counters[level] = n + 1
            return f'{n}. '
        return ['\u2022 ', '\u25E6 ', '\u25AA '][min(level, 2)]

    def _process_code(self, parent, tag):
        code_tag = tag.find('code')
        text = (code_tag.get_text() if code_tag else tag.get_text()).rstrip('\n')

        p = parent.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.0

        run = p.add_run(text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)

        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
        p.paragraph_format.element.get_or_add_pPr().append(shading)

    def _process_blockquote(self, parent, tag):
        for child in tag.children:
            if isinstance(child, NavigableString):
                continue
            if isinstance(child, Tag) and child.name == 'p':
                p = parent.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.right_indent = Inches(0.3)

                pPr = p.paragraph_format.element.get_or_add_pPr()
                pBdr = parse_xml(
                    f'<w:pBdr {nsdecls("w")}>'
                    f'  <w:left w:val="single" w:sz="12" w:space="8" w:color="C0C0C0"/>'
                    f'</w:pBdr>'
                )
                pPr.append(pBdr)

                fmt = {'italic': True}
                for c in child.children:
                    if isinstance(c, Tag):
                        self._process_inline(p, c, fmt)
                    elif isinstance(c, NavigableString):
                        self._emit_text(p, str(c), fmt)
            elif isinstance(child, Tag):
                self._process_children(parent, child)

    def _process_table(self, parent, tag):
        rows = tag.find_all('tr')
        if not rows:
            return

        has_header = bool(rows[0].find('th'))
        num_cols = max(len(r.find_all(['th', 'td'])) for r in rows)
        if num_cols == 0:
            return

        table = parent.add_table(rows=len(rows), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
            f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
            f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
            f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)

        for i, row_tag in enumerate(rows):
            cells = row_tag.find_all(['th', 'td'])
            for j, cell_tag in enumerate(cells):
                if j >= num_cols:
                    break
                cell = table.cell(i, j)
                cell.paragraphs[0].clear()

                is_h = cell_tag.name == 'th'
                if has_header and i == 0:
                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
                    cell.paragraphs[0].paragraph_format.element.get_or_add_pPr().append(shd)

                self._process_inline(cell.paragraphs[0], cell_tag, {'bold': is_h})

        p = parent.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
